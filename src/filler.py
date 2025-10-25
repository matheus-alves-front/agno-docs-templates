from docx import Document
from typing import Dict
import os
from .config import settings

def replace_runs(paragraph, mapping: Dict[str, str]):
    for i, run in enumerate(paragraph.runs):
        t = run.text or ""
        for k, v in mapping.items():
            if k in t:
                t = t.replace(k, v)
        run.text = t

def fill_docx(template_path: str, mapping: Dict[str, str], out_name: str | None = None) -> str:
    doc = Document(template_path)
    for p in doc.paragraphs:
        replace_runs(p, mapping)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_runs(p, mapping)
    if not out_name:
        base = os.path.splitext(os.path.basename(template_path))[0]
        out_name = f"{base}_preenchido.docx"
    out_path = os.path.join(settings.RESULTS, out_name)
    doc.save(out_path)
    return out_path
