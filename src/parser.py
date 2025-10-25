from typing import Dict, Any
from docx import Document
import re
import os
import unicodedata
from collections import OrderedDict
from .logging_utils import setup_logger, jlog

# Accept broader Unicode/ASCII placeholder names. Avoid greedy match ending at first closing brace.
PLACEHOLDER_RE = re.compile(r"\{([^{}:\s]+)(?::([^{}]+))?\}")
ENTITY_SUFFIXES = {
    "OUTORGANTE",
    "OUTORGADO",
    "COMPRADOR",
    "VENDEDOR",
    "DOADOR",
    "DONATARIO",
    "PF",
    "PJ",
    "DEVEDOR",
    "CREDOR",
    "OUTORGADA",
    "OUTORGANTES",
    "OUTORGADOS",
    "COMPRADORA",
    "VENDEDORA",
}

ENTITY_MIN_OCCURRENCES = 2
ENTITY_FIELD_HINTS = {
    "NOME",
    "CPF",
    "RG",
    "CNPJ",
    "ENDERECO",
    "LOGRADOURO",
    "BAIRRO",
    "CEP",
    "TELEFONE",
    "CELULAR",
    "WHATSAPP",
    "EMAIL",
    "REPRESENTANTE",
    "SOCIO",
    "SOCIA",
    "SOCIOS",
    "SOCIAS",
    "CONJUGE",
    "ESPOSO",
    "ESPOSA",
    "NACIONALIDADE",
    "NATURALIDADE",
    "PROFISSAO",
    "ESTADO",
    "UF",
    "MUNICIPIO",
    "CIDADE",
    "INSCRICAO",
    "RAZAO",
    "SOCIAL",
    "RAZAOSOCIAL",
    "INSCRICAOESTADUAL",
    "PODERES",
    "CAPITAL",
    "DATA",
    "NUMERO",
    "DOCUMENTO",
}


def _norm_token(token: str) -> str:
    return token.strip().casefold()


def _strip_accents(text: str) -> str:
    normalized = unicodedata.normalize("NFKD", text)
    return "".join(ch for ch in normalized if not unicodedata.combining(ch))


def read_docx_placeholders(path: str) -> Dict[str, Any]:
    doc = Document(path)
    texts = []

    for paragraph in doc.paragraphs:
        text = (paragraph.text or "").strip()
        if text:
            texts.append(text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = (cell.text or "").strip()
                if text:
                    texts.append(text)

    placeholders, titles, all_names = [], [], []
    for text in texts:
        for match in PLACEHOLDER_RE.finditer(text):
            name = match.group(1)
            placeholders.append({"raw": match.group(0), "name": name, "hint": match.group(2) or ""})
            all_names.append(name)
        upper = text.upper()
        if (text.endswith(":") or upper.isupper()) and len(text) < 120:
            titles.append(text.strip())

    return {"texts": texts, "placeholders": placeholders, "titles": titles, "all_names": all_names}


def infer_multiplicity_from_filename(fname: str) -> str:
    base = os.path.basename(fname)
    if "_V_V" in base:
        return "[V-V]"
    if "_V_1" in base:
        return "[V-1]"
    if "_1_V" in base:
        return "[1-V]"
    return "[1-1]"


def build_spec_from_docx(path: str) -> Dict[str, Any]:
    logger = setup_logger()
    filename = os.path.basename(path)
    jlog(logger, "INFO", "INDEX_START", file=filename)

    data = read_docx_placeholders(path)
    placeholders = data["placeholders"]
    all_names = list(dict.fromkeys(data["all_names"]))
    jlog(logger, "INFO", "INDEX_PLACEHOLDERS", file=filename, count=len(placeholders))

    entities = set()
    fields_by_key: "OrderedDict[tuple[str, str], Dict[str, Any]]" = OrderedDict()
    entity_max_indices: Dict[str, int] = {}
    # Pre-compute dynamic entity candidates based on placeholder suffix frequency.
    suffix_counts: Dict[str, int] = {}
    suffix_original: Dict[str, str] = {}
    suffix_hint_tokens: Dict[str, set[str]] = {}
    for placeholder in placeholders:
        name = placeholder["name"]
        parts = [part for part in name.split("_") if part]
        if not parts:
            continue
        candidate = None
        if parts[-1].isdigit() and len(parts) >= 2:
            candidate = parts[-2]
        else:
            candidate = parts[-1]
        if not candidate:
            continue
        norm = _norm_token(candidate)
        if not norm:
            continue
        suffix_counts[norm] = suffix_counts.get(norm, 0) + 1
        suffix_original.setdefault(norm, candidate)
        base_parts = parts[:-2] if parts[-1].isdigit() and len(parts) >= 2 else parts[:-1]
        if base_parts:
            hints = suffix_hint_tokens.setdefault(norm, set())
            for base_part in base_parts:
                normalized = _strip_accents(base_part).replace(" ", "").upper()
                if normalized:
                    hints.add(normalized)

    entity_lookup: Dict[str, str] = {suffix.casefold(): suffix for suffix in ENTITY_SUFFIXES}
    for norm, count in suffix_counts.items():
        if count < ENTITY_MIN_OCCURRENCES:
            continue
        if norm not in entity_lookup:
            hints = suffix_hint_tokens.get(norm, set())
            if not (hints and hints.intersection(ENTITY_FIELD_HINTS)):
                continue
            # Use uppercase to align with existing specs.
            entity_lookup[norm] = suffix_original[norm].upper()
    for placeholder in placeholders:
        name = placeholder["name"]
        parts = [part for part in name.split("_") if part]
        entity = None
        base = name
        number_idx = None
        candidate = None
        if len(parts) >= 2 and parts[-1].isdigit():
            candidate = parts[-2]
            number_idx = parts[-1]
        elif parts:
            candidate = parts[-1]

        if candidate:
            norm_candidate = _norm_token(candidate)
            lookup_entity = entity_lookup.get(norm_candidate)
            if lookup_entity:
                entity = lookup_entity
                if number_idx and len(parts) >= 2:
                    base_parts = parts[:-2]
                else:
                    base_parts = parts[:-1]
                if base_parts:
                    base = "_".join(base_parts)
                else:
                    base = candidate
        # Track known entities for grouping.
        if entity:
            entities.add(entity)

        key = (entity or "GLOBAL", base)
        if key not in fields_by_key:
            fields_by_key[key] = {
                "entity": entity or "GLOBAL",
                "name": base,
                "placeholder": placeholder["raw"],
            }

        if entity and number_idx:
            try:
                idx = int(number_idx)
            except ValueError:
                idx = 1
            if idx >= 1:
                entity_max_indices[entity] = max(entity_max_indices.get(entity, 1), idx)

    fields = list(fields_by_key.values())
    groups = []
    gid = 1
    for entity in sorted(x for x in entities if x != "GLOBAL"):
        entity_fields = [field for field in fields if field["entity"] == entity]
        if entity_fields:
            groups.append({"id": f"G{gid}", "label": f"Qualificação {entity}", "fields": entity_fields})
            gid += 1
    global_fields = [field for field in fields if field["entity"] == "GLOBAL"]
    if global_fields:
        groups.append({"id": f"G{gid}", "label": "Dados do Ato", "fields": global_fields})

    multiplicity = infer_multiplicity_from_filename(path)
    meta = {"casais": "_OU_CASAIS" in filename}
    if entity_max_indices:
        meta["inferred_counts"] = entity_max_indices
    spec = {
        "name": os.path.splitext(filename)[0],
        "source": filename,
        "multiplicity": multiplicity,
        "entities": sorted(list(entities)) or [],
        "groups": groups,
        "meta": meta,
        "all_placeholders": all_names,
    }
    entities_payload = list(spec["entities"])
    if global_fields:
        entities_payload.append("GLOBAL")
    jlog(logger, "INFO", "INDEX_ENTITIES", file=filename, entities=entities_payload)
    jlog(logger, "INFO", "INDEX_MULTIPLICITY", file=filename, multiplicity=multiplicity, casais=meta["casais"])

    return spec
